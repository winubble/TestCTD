"""
Author: Lu Lu 
Date: 2020/12/17                                                                                                                                                                                                   

Before Use: Make sure your computer is connected to the mtc-vpn for connecting the serial sever purpose.
Note: It may takes a while for running the code. The default sample time is 20. 

In this Version, I am adding another window and make the UI to be more user-friendly

If I shut down the computer when I run this program, it will show
ConnectionResetError: [WinError 10054] An existing connection was forcibly closed by the remote host.

Added the code deal with the AML Metrec X

"""

# ~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
# Import Libraries
from tkinter import * 
import tkinter as tk
import tkinter.font as tkFont               
import telnetlib
import getpass 
import sys
import time
import datetime
from datetime import datetime
import numpy as np
from matplotlib import pyplot as plt
import os
from tkinter.ttk import *
import matplotlib
matplotlib.use("TkAgg")
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
from matplotlib.figure import Figure
import pandas as pd
import docx
from tkinter import ttk
from time import gmtime, strftime
from tkinter import messagebox

# ~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
# Define the variables
HOST = "10.136.20.111"
testPort = ''
testID = ''
standardPort = "4001"
# ~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
# set up the main window
mainWin = tk.Tk()

mainWin.title("CTD Test Tool")
mainWin.geometry('625x400')

mainWin.configure(bg = "#737CA1")

fontStyle = tkFont.Font(family="Times New Roman", size = 15)

#mainWin.iconbitmap(r'sea.ico')

#~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~ greeting 
greeting = Label(mainWin, text = "Enter Test CTD Information Below:", font=fontStyle, background = "#737CA1", foreground = "white")
greeting.grid(row = 0, column = 1)
#~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~  IP 

askIP = Label(mainWin, text = "Host IP Address", font=fontStyle, background = "#737CA1", foreground = "white")
askIP.grid(row = 1, column = 0)

entryIP = Entry(mainWin, width = 15, cursor='star')
entryIP.insert(0,HOST)
entryIP.grid(row = 1, column = 1)
#~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~ Test Port number 

askTPort = Label(mainWin, text = "TEST Port Number", font=fontStyle, background = "#737CA1", foreground = "white")
askTPort.grid(row = 3, column = 0)

entryTPort = Entry(mainWin, width = 15, cursor='star')
entryTPort.grid(row = 3, column = 1)

#~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~ Standard Port number 

askSPort = Label(mainWin, text = "Standard Port Number", font=fontStyle, background = "#737CA1", foreground = "white")
askSPort.grid(row = 2, column = 0)

entrySPort = Entry(mainWin, width = 15, cursor='star')
entrySPort.insert(0,standardPort)
entrySPort.grid(row = 2, column = 1)

#~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~ Model Select
askModel = Label(mainWin, text = "CTD Model Select", font=fontStyle, background = "#737CA1", foreground = "white")
askModel.grid(row = 4, column = 0)
"""
# create an dropdown list for select the model
clicked = StringVar()
dropModel = OptionMenu(mainWin,clicked,"SBE-37", "SBE-19", "AML Metrec-X")
dropModel.grid(row = 3, column = 1)
"""
# create the radio button for get the model of the test CTD
# option 1 -> SBE-37
# option 2 -> SBE-19 
# option 3 -> AML Metrec-X

def findModel():
    if int(var.get()) == 1:
        #print("This is 37")
        myButton = Button(mainWin, text = "ENTER", command = get37)
        myButton.grid(row = 8, column = 3)
 
    if int(var.get()) == 2:
        #print("This is 19")
        myButton = Button(mainWin, text = "ENTER", command = get19)
        myButton.grid(row = 8, column = 3)

    if int(var.get()) == 3:
        #print("This is AML Metrec-X")
        myButton = Button(mainWin, text = "ENTER", command = getMetrecX)
        myButton.grid(row = 8, column = 3)

var = IntVar()

R1 = Radiobutton(mainWin, text="SBE-37", variable=var, value=1, command= findModel)
R1.grid(row = 4, column = 1)
R2 = Radiobutton(mainWin, text="SBE-19", variable=var, value=2, command= findModel)
R2.grid(row = 5, column = 1)
R3 = Radiobutton(mainWin, text="Metrec-X", variable=var, value=3, command= findModel)
R3.grid(row = 6, column = 1)
#~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~ sample progress

myBar = Label(mainWin, text = "Sampling Progress ", font=fontStyle, background = "#737CA1", foreground = "white")
myBar.grid(row = 7, column = 0, padx = 10, pady = 10)

# Add a progress bar
pbar = ttk.Progressbar(mainWin, length = 100, cursor='spider',
                     mode="determinate",
                     orient=tk.HORIZONTAL)
pbar.grid(row = 7, column = 1)

#=================================================================================================
#=================================================================================================
def get37():
    # validate the port number amd ip 
    testPort = entryTPort.get()
    standardPort = entrySPort.get()
    HOST = entryIP.get()

    try:
        # set up the conncetion with the test CTD
        tn2 = telnetlib.Telnet(HOST, testPort, 5)
        tn2.write(("\r\n").encode('ascii'))
        time.sleep(2)
    except:
        messagebox.showerror("Error", "Invalid Port Number or IP Address.")
        return
    try:
        # set up the conncetion with the standard CTD
        tn1 = telnetlib.Telnet(HOST, standardPort, 5)
        tn1.write(("\r\n").encode('ascii'))
        time.sleep(2)
    except:
        messagebox.showerror("Error", "Invalid Port Number or IP Address.")
        return
    
    # Ask the id for the test CTD
    tn2.write(("id?\r\n").encode('ascii'))
    time.sleep(5)
    
    # Get the testID
    test_idString = tn2.read_very_eager()

    id1 = 'id='
    id2 = '\r'

    # make the string into byte type class
    id1 = str.encode(id1)
    id2 = str.encode(id2)

    # Now testID is in bytes
    testID = test_idString[test_idString.find(id1)+3 : test_idString.find(id2)]

    # I want to make testID(byte) into string type
    encoding = 'utf-8'
    testID = testID.decode(encoding)

    # get the standardID
    # Ask the id for the test CTD
    tn1.write(("id?\r\n").encode('ascii'))
    time.sleep(5)
    standard_idString = tn1.read_very_eager()
    standardID = standard_idString[standard_idString.find(id1)+3 : standard_idString.find(id2)]
    standardID = standardID.decode(encoding)

    # send foramt command, I want them in XML format
    tn1.write(("#%sOutputFormat=2\r\n" %standardID).encode('ascii'))
    tn2.write(("#%sOutputFormat=2\r\n" %testID).encode('ascii'))
    time.sleep(5)

    # Set up arrays for storing the sample result 
    tempArr1 = []
    tempArr2 = []
    conArr1 = []
    conArr2 = []
    presArr1 = []
    presArr2 = []

    # Cumulate temperature
    tempCum = 0
    # Cumulate Conductivity
    conCum = 0
    # Pressure Decider
    isPres = 0

    # check if pressure exist 
    existPressure = False

    # Command for test CTD to take a sample
    testCommand = "#" + testID + "TS\r\n"
    standardCommand = "#" + standardID + "TS\r\n"

    # get the serial number of test CTD

    tn2.write((testCommand).encode('ascii'))
    time.sleep(10)
        
    # read and record the data
    data2 = str(tn2.read_very_eager())

    # get the serial number
    sn1 = '<sn>'
    sn2 = '</sn>'

    test_sn = str(data2[data2.find(sn1)+4 : data2.find(sn2)])

    # default take a sample 20 times 
    for i in range(0, 20):
        
        # Ask both benchmark CTD and Test CTD to take a sample
        tn1.write((standardCommand).encode('ascii'))
        tn2.write((testCommand).encode('ascii'))
        time.sleep(10)
        
        # read and record the data
        data1 = str(tn1.read_very_eager())
        data2 = str(tn2.read_very_eager())

        # get the temperature
        t1 = '<t1> '
        t2 = '</t1>'

        temp1 = float(data1[data1.find(t1)+5 : data1.find(t2)])
        temp2 = float(data2[data2.find(t1)+5 : data2.find(t2)])

        tempDiff = abs(temp1-temp2)
        
        tempPercent = tempDiff/temp1
        
        tempCum += tempPercent

        tempArr1.append(temp1)
        tempArr2.append(temp2)

        # get the conductivity
        c1 = '<c1>'
        c2 = '</c1>'

        con1 = float(data1[data1.find(c1)+5 : data1.find(c2)])
        con2 = float(data2[data2.find(c1)+5 : data2.find(c2)])

        conDiff = abs(con1-con2)
        
        conPercent = conDiff/con1
        
        conCum += conPercent

        conArr1.append(con1)
        conArr2.append(con2)

        # try to get the pressure
        p1 = '<p1>'
        p2 = '</p1>'
        if data2.find(p1)>0 :
            existPressure = True
            pres1 = float(data1[data1.find(p1)+5 : data1.find(p2)])
            pres2 = float(data2[data2.find(p1)+5 : data2.find(p2)])

            if (pres1<3 and pres1>0)  and (pres2<3 and pres2>0):
                isPres += 1

            presArr1.append(pres1)
            presArr2.append(pres2)

        # update the progress bar
        pbar["value"] = (i+1)*5
        mainWin.update()

    # plot the figure
    
    # make the array into np array
    tempArr1 = np.array(tempArr1)
    tempArr2 = np.array(tempArr2)
    conArr1 = np.array(conArr1)
    conArr2 = np.array(conArr2)
    
    # Get the x axis length
    myLen = tempArr1.size

    # axies for the temp figure
    x = np.arange(0, myLen)
    y1 = tempArr1
    y2 = tempArr2 
    y3 = conArr1
    y4 = conArr2
    
    tempPlot1 = {'Sample_Time': x,
         'Temp_Standard': y1
        }

    tempPlot2 = {'Sample_Time': x,
         'Temp_Test': y2
        }

    conPlot1 = {'Sample_Time': x,
         'Con_Standard': y3
        }

    conPlot2 = {'Sample_Time': x,
         'Con_Test': y4
        }   
    
    dfTemp1 = pd.DataFrame(tempPlot1,columns=['Sample_Time','Temp_Standard'])
    dfTemp2 = pd.DataFrame(tempPlot2,columns=['Sample_Time','Temp_Test'])

    dfCon1 = pd.DataFrame(conPlot1,columns=['Sample_Time','Con_Standard'])
    dfCon2 = pd.DataFrame(conPlot2,columns=['Sample_Time','Con_Test'])

    
    # Set a temp figure on the GUI
    figure1 = plt.Figure(figsize=(5,4), dpi=100)
    ax1 = figure1.add_subplot(111)
    #line1 = FigureCanvasTkAgg(figure1, mainWin)
    #line1.get_tk_widget().grid(row = 4, column = 1)


    # Set a con figure on the GUI
    figure2 = plt.Figure(figsize=(5,4), dpi=100)
    ax2 = figure2.add_subplot(111)
    #line2 = FigureCanvasTkAgg(figure2, mainWin)
    #line2.get_tk_widget().grid(row = 4, column = 2)
    
    # plot the figure on the GUI 
    dfTemp1 = dfTemp1[['Sample_Time','Temp_Standard']].groupby('Sample_Time').sum()
    dfTemp2 = dfTemp2[['Sample_Time','Temp_Test']].groupby('Sample_Time').sum()

    dfCon1 = dfCon1[['Sample_Time','Con_Standard']].groupby('Sample_Time').sum()
    dfCon2 = dfCon2[['Sample_Time','Con_Test']].groupby('Sample_Time').sum()

    dfTemp1.plot(kind='line', legend=True, ax=ax1, marker= "D")
    dfTemp2.plot(kind='line', legend=True, ax=ax1, marker= "D")
    
    ax1.set_title('Temperature of two CTDs')

    dfCon1.plot(kind='line', legend=True, ax=ax2, marker= "D")
    dfCon2.plot(kind='line', legend=True, ax=ax2, marker= "D")

    ax2.set_title('Conductivity of two CTDs')

    divider = Label(mainWin, text = "--------------------",background = "#737CA1",  font=fontStyle, foreground = "white" )
    divider.grid(row = 9, column = 0)
    # deal with the pressure
    if existPressure == True:
        presArr1 = np.array(presArr1)
        presArr2 = np.array(presArr2)

        y5 = presArr1
        y6 = presArr2

        presPlot1 = {'Sample_Time': x,
         'Pres_Standard': y5
        }

        presPlot2 = {'Sample_Time': x,
         'Pres_Test': y6
        }

        dfPres1 = pd.DataFrame(presPlot1,columns=['Sample_Time','Pres_Standard'])
        dfPres2 = pd.DataFrame(presPlot2,columns=['Sample_Time','Pres_Test'])

        # Set a pressure figure on the GUI
        figure3 = plt.Figure(figsize=(5,4), dpi=100)
        ax3 = figure3.add_subplot(111)
        #line3 = FigureCanvasTkAgg(figure3, mainWin)
        #line3.get_tk_widget().grid(row = 4, column = 3)

        # plot the pressure figure
        dfPres1 = dfPres1[['Sample_Time','Pres_Standard']].groupby('Sample_Time').sum()
        dfPres2 = dfPres2[['Sample_Time','Pres_Test']].groupby('Sample_Time').sum()

        dfPres1.plot(kind='line', legend=True, ax=ax3, marker= "D")
        dfPres2.plot(kind='line', legend=True, ax=ax3, marker= "D")
    
        ax3.set_title('Pressure of two CTDs')

        # save the plot 
        figure3.savefig('presplot.jpg', bbox_inches='tight', dpi=150)

         # Check if pass the pressure test
        if (isPres == 20):
            ifPres = Label(mainWin, text = "Pressure Test Passed.", background = "#737CA1",  font=fontStyle, foreground = "white")
            ifPres.grid(row = 10, column = 0)
        else:
            ifPres = Label(mainWin, text = "Pressure Test Failed.", background = "#737CA1",  font=fontStyle, foreground = "white")
            ifPres.grid(row = 10, column = 0)

    # save the plot
    figure1.savefig('tempplot.jpg', bbox_inches='tight', dpi=150)
    figure2.savefig('conplot.jpg', bbox_inches='tight', dpi=150)
    
    # Check if pass the temp test
    if (tempCum/myLen <= 0.01):
        ifTemp = Label(mainWin, text = "Temperature Test Passed.", background = "#737CA1",  font=fontStyle, foreground = "white")
        ifTemp.grid(row = 11, column = 0)
    else:
        ifTemp = Label(mainWin, text = "Temperature Test Failed.", background = "#737CA1",  font=fontStyle, foreground = "white")
        ifTemp.grid(row = 11, column = 0)

    # Check if pass the con test
    if (conCum/myLen <= 0.01):
        ifCon = Label(mainWin, text = "Conductivity Test Passed.", background = "#737CA1",  font=fontStyle, foreground = "white")
        ifCon.grid(row = 12, column = 0)
    else:
        ifCon = Label(mainWin, text = "Conductivity Test Failed.", background = "#737CA1",  font=fontStyle, foreground = "white")
        ifCon.grid(row = 12, column = 0)
    
    # Write the doc report  
    timeString = time.ctime()
    myReport = docx.Document()
    timeZone = strftime("%Z")
    myReport.add_heading("CTD Test Report", 0)
    myReport.add_paragraph("This is a test report for testing CTD of port number "+ str(testPort) + " on "+ timeString + " of "+ timeZone + ".")
    myReport.add_paragraph("The Model of Test CTD is SBE-37 and its ID is " + str(testID) + ".")
    myReport.add_paragraph("The serial Server IP Address is "+ str(HOST) + ". The sample time is 20.")
    myReport.add_paragraph("The report contains three parts: Test Data, Test Figure and Test Result.")

    #~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
    myReport.add_heading("Test Data", 1)
    
    if existPressure == False: 
        table = myReport.add_table(rows=1, cols=5)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Sample Time'
        hdr_cells[1].text = 'Standard Temperature'
        hdr_cells[2].text = 'Test Temperature'
        hdr_cells[3].text = 'Standard Conductivity'
        hdr_cells[4].text = 'Test Conductivity'

        for i in range(0, 20):
            row_cells = table.add_row().cells
            row_cells[0].text = str(i+1)
            row_cells[1].text = str(tempArr1[i])
            row_cells[2].text = str(tempArr2[i])
            row_cells[3].text = str(conArr1[i])
            row_cells[4].text = str(conArr2[i])
    else:
        table = myReport.add_table(rows=1, cols=7)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Sample Time'
        hdr_cells[1].text = 'Standard Temperature'
        hdr_cells[2].text = 'Test Temperature'
        hdr_cells[3].text = 'Standard Conductivity'
        hdr_cells[4].text = 'Test Conductivity'
        hdr_cells[5].text = 'Standard Pressure'
        hdr_cells[6].text = 'Test Pressure'

        for i in range(0, 20):
            row_cells = table.add_row().cells
            row_cells[0].text = str(i+1)
            row_cells[1].text = str(tempArr1[i])
            row_cells[2].text = str(tempArr2[i])
            row_cells[3].text = str(conArr1[i])
            row_cells[4].text = str(conArr2[i])
            row_cells[5].text = str(presArr1[i])
            row_cells[6].text = str(presArr2[i])

#~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~

    myReport.add_heading("Test Figure", 1)
    myReport.add_paragraph("Temperature Figure: ")
    myReport.add_picture("tempplot.jpg", width=docx.shared.Inches(4), height=docx.shared.Inches(4))
    myReport.add_paragraph("Conductivity Figure: ")
    myReport.add_picture("conplot.jpg", width=docx.shared.Inches(4), height=docx.shared.Inches(4))

    if existPressure == True:
        myReport.add_paragraph("Pressure Figure: ")
        myReport.add_picture("presplot.jpg", width=docx.shared.Inches(4), height=docx.shared.Inches(4))
#~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~

    myReport.add_heading("Test Result ", 1)
    myReport.add_paragraph("Note: If the difference of test CTD's data and standard CTD's data is within 1%, then the corresponding test is passed. If all the Temperature, Conductivity and Pressure(if exists) Tests passed, then final result will be passed.")

    final_pass_or_fail = False

    if (tempCum/myLen <= 0.01):
        myReport.add_paragraph("Temperature Test: Passed")
    else:
        myReport.add_paragraph("Temperature Test: Failed")

    if (conCum/myLen <= 0.01):
        myReport.add_paragraph("Conductivity Test: Passed")
    else:
        myReport.add_paragraph("Conductivity Test: Failed")

    if existPressure == True:
        if (isPres == 20):
            myReport.add_paragraph("Pressure Test: Passed")
        else:
            myReport.add_paragraph("Pressure Test: Failed")

        if (tempCum/myLen <= 0.01) and (conCum/myLen <= 0.01) and (isPres == 20):
            final_pass_or_fail = True
            final = myReport.add_paragraph("Final Result: ")
            final.add_run('Pass!').bold = True
        else:
            final = myReport.add_paragraph("Final Result: ")
            final.add_run('Fail!').bold = True
    else:
        if (tempCum/myLen <= 0.01) and (conCum/myLen <= 0.01):
            final_pass_or_fail = True
            final = myReport.add_paragraph("Final Result: ")
            final.add_run('Pass!').bold = True
        else:
            final = myReport.add_paragraph("Final Result: ")
            final.add_run('Fail!').bold = True

    #Title report in format "yyyy-mm-dd_SBE##_SN:xxxx_PASS" where ## = model(ie 37 or 19) and xxxx = serial number 
    if final_pass_or_fail:
        reportResult = "PASS"
    else:
        reportResult = "FAIL"

    reportTime = datetime.today().strftime('%Y-%m-%d')
    reportName = "%s_SBE37_SN%s_%s.docx" % (reportTime, test_sn, reportResult)
    try:
        myReport.save(r'%s' % reportName)

        # end words on GUI 
        ending = Label(mainWin, text = "Report successfully generated.", background = "#737CA1",  font=fontStyle, foreground = "white")
        ending.grid(row = 13, column = 0)
    except:
        ending = Label(mainWin, text = "Fail to generate the report.", background = "#737CA1",  font=fontStyle, foreground = "white")
        ending.grid(row = 13, column = 0)
        return

    tn1.close()
    tn2.close()
#=================================================================================================
#=================================================================================================
def get19():
    # validate the port number amd ip 
    testPort = entryTPort.get()
    standardPort = entrySPort.get()
    HOST = entryIP.get()

    try:
        # set up the conncetion with the test CTD
        tn2 = telnetlib.Telnet(HOST, testPort, 5)
        tn2.write(("\r\n").encode('ascii'))
        time.sleep(2)
    except:
        messagebox.showerror("Error", "Invalid Port Number or IP Address.")
        return
    try:
        # set up the conncetion with the standard CTD
        tn1 = telnetlib.Telnet(HOST, standardPort, 5)
        tn1.write(("\r\n").encode('ascii'))
        time.sleep(2)
    except:
        messagebox.showerror("Error", "Invalid Port Number or IP Address.")
        return
    
    """For the SBE-19 model, we don't need to get the ID for the test CTD
    """

    # send foramt command, I want them in XML format
    # tn1 is the connection with standard CTD
    # tn2 is the conncetion with test CTD
    tn1.write(("#35OutputFormat=2\r\n").encode('ascii'))
    tn2.write(("OutputFormat=5\r\n").encode('ascii'))
    time.sleep(5)

    # Set up arrays for storing the sample result 
    tempArr1 = []
    tempArr2 = []
    conArr1 = []
    conArr2 = []
    presArr1 = []
    presArr2 = []

    # Cumulate temperature
    tempCum = 0
    # Cumulate Conductivity
    conCum = 0
    # Cumulate Pressure
    isPres = 0

    # check if pressure exist 
    existPressure = False

    # Command for test CTD to take a sample
    testCommand = "TS\r\n"

    # get the serial number of test CTD
    tn2.write((testCommand).encode('ascii'))
    time.sleep(10)
        
    # read and record the data
    data2 = str(tn2.read_very_eager())

    # get the serial number
    sn1 = '<sn>'
    sn2 = '</sn>'

    test_sn = str(data2[data2.find(sn1)+4 : data2.find(sn2)])

    # default take a sample 20 times 
    for i in range(0, 20):
        
        # Ask both benchmark CTD and Test CTD to take a sample
        tn1.write(("#35TS\r\n").encode('ascii'))
        tn2.write((testCommand).encode('ascii'))
        time.sleep(10)
        
        # read and record the data
        data1 = str(tn1.read_very_eager())
        data2 = str(tn2.read_very_eager())
        #print(data2)
        #print("\r")

        # get the temperature
        t1 = '<t1> '
        t2 = '</t1>'

        temp1 = float(data1[data1.find(t1)+5 : data1.find(t2)])
        temp2 = float(data2[data2.find(t1)+4 : data2.find(t2)])

        tempDiff = abs(temp1-temp2)
        
        tempPercent = tempDiff/temp1
        
        tempCum += tempPercent

        tempArr1.append(temp1)
        tempArr2.append(temp2)

        # get the conductivity
        c1 = '<c1>'
        c2 = '</c1>'

        con1 = float(data1[data1.find(c1)+5 : data1.find(c2)])
        con2 = float(data2[data2.find(c1)+4 : data2.find(c2)])

        conDiff = abs(con1-con2)
        
        conPercent = conDiff/con1
        
        conCum += conPercent

        conArr1.append(con1)
        conArr2.append(con2)

        # try to get the pressure
        p1 = '<p1>'
        p2 = '</p1>'
        if data2.find(p1)>0 :
            existPressure = True
            pres1 = float(data1[data1.find(p1)+5 : data1.find(p2)])
            pres2 = float(data2[data2.find(p1)+4 : data2.find(p2)])

            if (pres1<3 and pres1>0)  and (pres2<3 and pres2>0):
                isPres += 1

            presArr1.append(pres1)
            presArr2.append(pres2)

        # update the progress bar
        pbar["value"] = (i+1)*5
        mainWin.update()

    # plot the figure

    # make the array into np array
    tempArr1 = np.array(tempArr1)
    tempArr2 = np.array(tempArr2)
    conArr1 = np.array(conArr1)
    conArr2 = np.array(conArr2)
    
    # Get the x axis length
    myLen = tempArr1.size

    # axies for the temp figure
    x = np.arange(0, myLen)
    y1 = tempArr1
    y2 = tempArr2 
    y3 = conArr1
    y4 = conArr2
    
    tempPlot1 = {'Sample_Time': x,
         'Temp_Standard': y1
        }

    tempPlot2 = {'Sample_Time': x,
         'Temp_Test': y2
        }

    conPlot1 = {'Sample_Time': x,
         'Con_Standard': y3
        }

    conPlot2 = {'Sample_Time': x,
         'Con_Test': y4
        }
    
    
    dfTemp1 = pd.DataFrame(tempPlot1,columns=['Sample_Time','Temp_Standard'])
    dfTemp2 = pd.DataFrame(tempPlot2,columns=['Sample_Time','Temp_Test'])

    dfCon1 = pd.DataFrame(conPlot1,columns=['Sample_Time','Con_Standard'])
    dfCon2 = pd.DataFrame(conPlot2,columns=['Sample_Time','Con_Test'])

    # Set a temp figure on the GUI
    figure1 = plt.Figure(figsize=(5,4), dpi=100)
    ax1 = figure1.add_subplot(111)
    #line1 = FigureCanvasTkAgg(figure1, mainWin)
    #line1.get_tk_widget().grid(row = 4, column = 1)

    # Set a con figure on the GUI
    figure2 = plt.Figure(figsize=(5,4), dpi=100)
    ax2 = figure2.add_subplot(111)
    #line2 = FigureCanvasTkAgg(figure2, mainWin)
    #line2.get_tk_widget().grid(row = 4, column = 2)
    
    # plot the figure on the GUI 
    dfTemp1 = dfTemp1[['Sample_Time','Temp_Standard']].groupby('Sample_Time').sum()
    dfTemp2 = dfTemp2[['Sample_Time','Temp_Test']].groupby('Sample_Time').sum()

    dfCon1 = dfCon1[['Sample_Time','Con_Standard']].groupby('Sample_Time').sum()
    dfCon2 = dfCon2[['Sample_Time','Con_Test']].groupby('Sample_Time').sum()

    dfTemp1.plot(kind='line', legend=True, ax=ax1, marker= "D")
    dfTemp2.plot(kind='line', legend=True, ax=ax1, marker= "D")
    
    ax1.set_title('Temperature of two CTDs')

    dfCon1.plot(kind='line', legend=True, ax=ax2, marker= "D")
    dfCon2.plot(kind='line', legend=True, ax=ax2, marker= "D")

    ax2.set_title('Conductivity of two CTDs')

    # I am a cute divider
    divider = Label(mainWin, text = "--------------------",background = "#737CA1",  font=fontStyle, foreground = "white" )
    divider.grid(row = 9, column = 0)

    # deal with the pressure
    if existPressure == True:
        presArr1 = np.array(presArr1)
        presArr2 = np.array(presArr2)

        y5 = presArr1
        y6 = presArr2

        presPlot1 = {'Sample_Time': x,
         'Pres_Standard': y5
        }

        presPlot2 = {'Sample_Time': x,
         'Pres_Test': y6
        }

        dfPres1 = pd.DataFrame(presPlot1,columns=['Sample_Time','Pres_Standard'])
        dfPres2 = pd.DataFrame(presPlot2,columns=['Sample_Time','Pres_Test'])

        # Set a pressure figure on the GUI
        figure3 = plt.Figure(figsize=(5,4), dpi=100)
        ax3 = figure3.add_subplot(111)
        #line3 = FigureCanvasTkAgg(figure3, mainWin)
        #line3.get_tk_widget().grid(row = 4, column = 3)

        # plot the pressure figure
        dfPres1 = dfPres1[['Sample_Time','Pres_Standard']].groupby('Sample_Time').sum()
        dfPres2 = dfPres2[['Sample_Time','Pres_Test']].groupby('Sample_Time').sum()

        dfPres1.plot(kind='line', legend=True, ax=ax3, marker= "D")
        dfPres2.plot(kind='line', legend=True, ax=ax3, marker= "D")
    
        ax3.set_title('Pressure of two CTDs')

        # save the plot 
        figure3.savefig('presplot.jpg', bbox_inches='tight', dpi=150)

         # Check if pass the pressure test
        if (isPres == 20):
            ifPres = Label(mainWin, text = "Pressure Test Passed", background = "#737CA1",  font=fontStyle, foreground = "white")
            ifPres.grid(row = 10, column = 0)
        else:
            ifPres = Label(mainWin, text = "Pressure Test Failed", background = "#737CA1",  font=fontStyle, foreground = "white")
            ifPres.grid(row = 10, column = 0)

    # save the plot
    figure1.savefig('tempplot.jpg', bbox_inches='tight', dpi=150)
    figure2.savefig('conplot.jpg', bbox_inches='tight', dpi=150)
    
    # Check if pass the temp test
    if (tempCum/myLen <= 0.01):
        ifTemp = Label(mainWin, text = "Temperature Test Passed", background = "#737CA1",  font=fontStyle, foreground = "white")
        ifTemp.grid(row = 11, column = 0)
    else:
        ifTemp = Label(mainWin, text = "Temperature Test Failed", background = "#737CA1",  font=fontStyle, foreground = "white")
        ifTemp.grid(row = 11, column = 0)

    # Check if pass the con test
    if (conCum/myLen <= 0.01):
        ifCon = Label(mainWin, text = "Conductivity Test Passed", background = "#737CA1",  font=fontStyle, foreground = "white")
        ifCon.grid(row = 12, column = 0)
    else:
        ifCon = Label(mainWin, text = "Conductivity Test Failed", background = "#737CA1",  font=fontStyle, foreground = "white")
        ifCon.grid(row = 12, column = 0)

     # Write the doc report  
    timeString = time.ctime()
    myReport = docx.Document()
    timeZone = strftime("%Z")
    myReport.add_heading("CTD Test Report", 0)
    myReport.add_paragraph("This is a test report for testing CTD of port number "+ str(testPort) + " on "+ timeString + " of "+ timeZone + ".")
    myReport.add_paragraph("The Model of Test CTD is SBE-19.")
    myReport.add_paragraph("The serial Server IP Address is "+ str(HOST) + ". The sample time is 20.")
    myReport.add_paragraph("The report contains three parts: Test Data, Test Figure and Test Result.")

    #~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
    myReport.add_heading("Test Data", 1)
    
    if existPressure == False: 
        table = myReport.add_table(rows=1, cols=5)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Sample Time'
        hdr_cells[1].text = 'Standard Temperature'
        hdr_cells[2].text = 'Test Temperature'
        hdr_cells[3].text = 'Standard Conductivity'
        hdr_cells[4].text = 'Test Conductivity'

        for i in range(0, 20):
            row_cells = table.add_row().cells
            row_cells[0].text = str(i+1)
            row_cells[1].text = str(tempArr1[i])
            row_cells[2].text = str(tempArr2[i])
            row_cells[3].text = str(conArr1[i])
            row_cells[4].text = str(conArr2[i])
    else:
        table = myReport.add_table(rows=1, cols=7)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Sample Time'
        hdr_cells[1].text = 'Standard Temperature'
        hdr_cells[2].text = 'Test Temperature'
        hdr_cells[3].text = 'Standard Conductivity'
        hdr_cells[4].text = 'Test Conductivity'
        hdr_cells[5].text = 'Standard Pressure'
        hdr_cells[6].text = 'Test Pressure'

        for i in range(0, 20):
            row_cells = table.add_row().cells
            row_cells[0].text = str(i+1)
            row_cells[1].text = str(tempArr1[i])
            row_cells[2].text = str(tempArr2[i])
            row_cells[3].text = str(conArr1[i])
            row_cells[4].text = str(conArr2[i])
            row_cells[5].text = str(presArr1[i])
            row_cells[6].text = str(presArr2[i])

#~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~

    myReport.add_heading("Test Figure", 1)
    myReport.add_paragraph("Temperature Figure: ")
    myReport.add_picture("tempplot.jpg", width=docx.shared.Inches(4), height=docx.shared.Inches(4))
    myReport.add_paragraph("Conductivity Figure: ")
    myReport.add_picture("conplot.jpg", width=docx.shared.Inches(4), height=docx.shared.Inches(4))

    if existPressure == True:
        myReport.add_paragraph("Pressure Figure: ")
        myReport.add_picture("presplot.jpg", width=docx.shared.Inches(4), height=docx.shared.Inches(4))
#~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~

    myReport.add_heading("Test Result ", 1)
    myReport.add_paragraph("Note: If the difference of test CTD's data and standard CTD's data is within 1%, then the corresponding test is passed. If all the Temperature, Conductivity and Pressure(if exists) Tests passed, then final result will be passed.")

    final_pass_or_fail = False

    if (tempCum/myLen <= 0.01):
        myReport.add_paragraph("Temperature Test: Passed")
    else:
        myReport.add_paragraph("Temperature Test: Failed")

    if (conCum/myLen <= 0.01):
        myReport.add_paragraph("Conductivity Test: Passed")
    else:
        myReport.add_paragraph("Conductivity Test: Failed")

    if existPressure == True:
        if (isPres == 20):
            myReport.add_paragraph("Pressure Test: Passed")
        else:
            myReport.add_paragraph("Pressure Test: Failed")

        if (tempCum/myLen <= 0.01) and (conCum/myLen <= 0.01) and (isPres == 20):
            final_pass_or_fail = True
            final = myReport.add_paragraph("Final Result: ")
            final.add_run('Pass!').bold = True
        else:
            final = myReport.add_paragraph("Final Result: ")
            final.add_run('Fail!').bold = True
    else:
        if (tempCum/myLen <= 0.01) and (conCum/myLen <= 0.01):
            final_pass_or_fail = True
            final = myReport.add_paragraph("Final Result: ")
            final.add_run('Pass!').bold = True
        else:
            final = myReport.add_paragraph("Final Result: ")
            final.add_run('Fail!').bold = True

    #Title report in format "yyyy-mm-dd_SBE##_SN:xxxx_PASS" where ## = model(ie 37 or 19) and xxxx = serial number 
    if final_pass_or_fail:
        reportResult = "PASS"
    else:
        reportResult = "FAIL"

    reportTime = datetime.today().strftime('%Y-%m-%d')
    reportName = "%s_SBE19_SN%s_%s.docx" % (reportTime, test_sn, reportResult)

    try:
        myReport.save(r'%s' % reportName)
        # end words on GUI 
        ending = Label(mainWin, text = "Report successfully generated.", background = "#737CA1",  font=fontStyle, foreground = "white")
        ending.grid(row = 13, column = 0)
    except:
        ending = Label(mainWin, text = "Fail to generate the report.", background = "#737CA1",  font=fontStyle, foreground = "white")
        ending.grid(row = 13, column = 0)
        return

    tn1.close()
    tn2.close()
#=================================================================================================
#=================================================================================================

def getMetrecX():
    # validate the port number amd ip 
    testPort = entryTPort.get()
    standardPort = entrySPort.get()
    HOST = entryIP.get()

    try:
        # set up the conncetion with the test CTD
        tn2 = telnetlib.Telnet(HOST, testPort, 5)
        tn2.write(("\r\n").encode('ascii'))
        time.sleep(2)
    except:
        messagebox.showerror("Error", "Invalid Port Number or IP Address.")
        return
    try:
        # set up the conncetion with the standard CTD
        tn1 = telnetlib.Telnet(HOST, standardPort, 5)
        tn1.write(("\r\n").encode('ascii'))
        time.sleep(2)
    except:
        messagebox.showerror("Error", "Invalid Port Number or IP Address.")
        return
    
    # Change the standard CTD output to XML 
    tn1.write(("#35OutputFormat=2\r\n").encode('ascii'))
    time.sleep(5)

    # Set up arrays for storing the sample result 
    tempArr1 = []
    tempArr2 = []
    conArr1 = []
    conArr2 = []
    presArr1 = []
    presArr2 = []

    # Cumulate temperature
    tempCum = 0
    # Cumulate Conductivity
    conCum = 0
    # Check Pressure
    isPres = 0

    # get the serial number of test CTD
    tn2.write(("version\r\n").encode('ascii'))
    time.sleep(10)
        
    # read and record the data
    data2 = str(tn2.read_very_eager())

    # get the serial number
    sn1 = 'SN:'
    test_sn = str(data2[data2.find(sn1)+3 : data2.find(sn1)+8])

    #skip the first record since the format is not the same
    tn2.write(("scan\r\n").encode('ascii'))
    time.sleep(10)
    tn2.read_very_eager()

    for i in range(0, 20):
        
        # Ask both benchmark CTD and Test CTD to take a sample
        tn1.write(("#35TS\r\n").encode('ascii'))
        tn2.write(("scan\r\n").encode('ascii'))
        time.sleep(10)
        
        # read and record the data
        data1 = str(tn1.read_very_eager())
        data2 = str(tn2.read_very_eager())
        #print(data2)

        # get the temperature
        t1 = '<t1> '
        t2 = '</t1>'

        temp1 = float(data1[data1.find(t1)+5 : data1.find(t2)])
        temp2 = float(data2[42:48])
        #print(temp2)

        tempDiff = abs(temp1-temp2)
        
        tempPercent = tempDiff/temp1
        
        tempCum += tempPercent

        tempArr1.append(temp1)
        tempArr2.append(temp2)

        # get the conductivity
        c1 = '<c1>'
        c2 = '</c1>'

        con1 = float(data1[data1.find(c1)+5 : data1.find(c2)])
        con2 = (float(data2[34:40]))/10
        #print(con2)

        conDiff = abs(con1-con2)
        
        conPercent = conDiff/con1
        
        conCum += conPercent

        conArr1.append(con1)
        conArr2.append(con2)

        # get the pressure
        p1 = '<p1>'
        p2 = '</p1>'
        
        pres1 = float(data1[data1.find(p1)+5 : data1.find(p2)])
        pres2 = float(data2[50:58])
        #print(pres2)

        if (pres1<3 and pres1>0)  and (pres2<3 and pres2>0):
            isPres += 1

        presArr1.append(pres1)
        presArr2.append(pres2)

        # update the progress bar
        pbar["value"] = (i+1)*5
        mainWin.update()

    # plot the figure
    # make the array into np array
    tempArr1 = np.array(tempArr1)
    tempArr2 = np.array(tempArr2)
    conArr1 = np.array(conArr1)
    conArr2 = np.array(conArr2)
    presArr1 = np.array(presArr1)
    presArr2 = np.array(presArr2)
    
    # Get the x axis length
    myLen = tempArr1.size

    # axies for the temp figure
    x = np.arange(0, myLen)
    y1 = tempArr1
    y2 = tempArr2 
    y3 = conArr1
    y4 = conArr2
    y5 = presArr1
    y6 = presArr2
    
    tempPlot1 = {'Sample_Time': x,
         'Temp_Standard': y1
        }

    tempPlot2 = {'Sample_Time': x,
         'Temp_Test': y2
        }

    conPlot1 = {'Sample_Time': x,
         'Con_Standard': y3
        }

    conPlot2 = {'Sample_Time': x,
         'Con_Test': y4
        }
    presPlot1 = {'Sample_Time': x,
         'Pres_Standard': y5
        }

    presPlot2 = {'Sample_Time': x,
         'Pres_Test': y6
    }
    
    
    dfTemp1 = pd.DataFrame(tempPlot1,columns=['Sample_Time','Temp_Standard'])
    dfTemp2 = pd.DataFrame(tempPlot2,columns=['Sample_Time','Temp_Test'])

    dfCon1 = pd.DataFrame(conPlot1,columns=['Sample_Time','Con_Standard'])
    dfCon2 = pd.DataFrame(conPlot2,columns=['Sample_Time','Con_Test'])

    dfPres1 = pd.DataFrame(presPlot1,columns=['Sample_Time','Pres_Standard'])
    dfPres2 = pd.DataFrame(presPlot2,columns=['Sample_Time','Pres_Test'])

    # Set a temp figure on the GUI
    figure1 = plt.Figure(figsize=(5,4), dpi=100)
    ax1 = figure1.add_subplot(111)
    #line1 = FigureCanvasTkAgg(figure1, mainWin)
    #line1.get_tk_widget().grid(row = 4, column = 1)

    # Set a con figure on the GUI
    figure2 = plt.Figure(figsize=(5,4), dpi=100)
    ax2 = figure2.add_subplot(111)
    #line2 = FigureCanvasTkAgg(figure2, mainWin)
    #line2.get_tk_widget().grid(row = 4, column = 2)

    # Set a pressure figure on the GUI
    figure3 = plt.Figure(figsize=(5,4), dpi=100)
    ax3 = figure3.add_subplot(111)
    #line3 = FigureCanvasTkAgg(figure3, mainWin)
    #line3.get_tk_widget().grid(row = 4, column = 3)
    
    # plot the figure 
    dfTemp1 = dfTemp1[['Sample_Time','Temp_Standard']].groupby('Sample_Time').sum()
    dfTemp2 = dfTemp2[['Sample_Time','Temp_Test']].groupby('Sample_Time').sum()

    dfCon1 = dfCon1[['Sample_Time','Con_Standard']].groupby('Sample_Time').sum()
    dfCon2 = dfCon2[['Sample_Time','Con_Test']].groupby('Sample_Time').sum()

    # plot the pressure figure
    dfPres1 = dfPres1[['Sample_Time','Pres_Standard']].groupby('Sample_Time').sum()
    dfPres2 = dfPres2[['Sample_Time','Pres_Test']].groupby('Sample_Time').sum()


    dfTemp1.plot(kind='line', legend=True, ax=ax1, marker= "D")
    dfTemp2.plot(kind='line', legend=True, ax=ax1, marker= "D")
    
    ax1.set_title('Temperature of two CTDs')

    dfCon1.plot(kind='line', legend=True, ax=ax2, marker= "D")
    dfCon2.plot(kind='line', legend=True, ax=ax2, marker= "D")

    ax2.set_title('Conductivity of two CTDs')

    dfPres1.plot(kind='line', legend=True, ax=ax3, marker= "D")
    dfPres2.plot(kind='line', legend=True, ax=ax3, marker= "D")
    
    ax3.set_title('Pressure of two CTDs')

    # save the plot
    figure1.savefig('tempplot.jpg', bbox_inches='tight', dpi=150)
    figure2.savefig('conplot.jpg', bbox_inches='tight', dpi=150)
    figure3.savefig('presplot.jpg', bbox_inches='tight', dpi=150)

    # I am a cute divider
    divider = Label(mainWin, text = "--------------------",background = "#737CA1",  font=fontStyle, foreground = "white" )
    divider.grid(row = 9, column = 0)

    # Check if pass the temp test
    if (tempCum/myLen <= 0.01):
        ifTemp = Label(mainWin, text = "Temperature Test Passed", background = "#737CA1",  font=fontStyle, foreground = "white")
        ifTemp.grid(row = 11, column = 0)
    else:
        ifTemp = Label(mainWin, text = "Temperature Test Failed", background = "#737CA1",  font=fontStyle, foreground = "white")
        ifTemp.grid(row = 11, column = 0)

    # Check if pass the con test
    if (conCum/myLen <= 0.01):
        ifCon = Label(mainWin, text = "Conductivity Test Passed", background = "#737CA1",  font=fontStyle, foreground = "white")
        ifCon.grid(row = 12, column = 0)
    else:
        ifCon = Label(mainWin, text = "Conductivity Test Failed", background = "#737CA1",  font=fontStyle, foreground = "white")
        ifCon.grid(row = 12, column = 0)

    # Check if pass the pressure test
    if (isPres == 20):
        ifPres = Label(mainWin, text = "Pressure Test Passed", background = "#737CA1",  font=fontStyle, foreground = "white")
        ifPres.grid(row = 10, column = 0)
    else:
        ifPres = Label(mainWin, text = "Pressure Test Failed", background = "#737CA1",  font=fontStyle, foreground = "white")
        ifPres.grid(row = 10, column = 0)

    # Write the doc report  
    timeString = time.ctime()
    myReport = docx.Document()
    timeZone = strftime("%Z")
    myReport.add_heading("CTD Test Report", 0)
    myReport.add_paragraph("This is a test report for testing CTD of port number "+ str(testPort) + " on "+ timeString + " of "+ timeZone + ".")
    myReport.add_paragraph("The Model of Test CTD is AML Metrec-X.")
    myReport.add_paragraph("The serial Server IP Address is "+ str(HOST) + ". The sample time is 20.")
    myReport.add_paragraph("The report contains three parts: Test Data, Test Figure and Test Result.")  

    #~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
    myReport.add_heading("Test Data", 1)
    
    table = myReport.add_table(rows=1, cols=7)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Sample Time'
    hdr_cells[1].text = 'Standard Temperature'
    hdr_cells[2].text = 'Test Temperature'
    hdr_cells[3].text = 'Standard Conductivity'
    hdr_cells[4].text = 'Test Conductivity'
    hdr_cells[5].text = 'Standard Pressure'
    hdr_cells[6].text = 'Test Pressure'

    for i in range(0, 20):
        row_cells = table.add_row().cells
        row_cells[0].text = str(i+1)
        row_cells[1].text = str(tempArr1[i])
        row_cells[2].text = str(tempArr2[i])
        row_cells[3].text = str(conArr1[i])
        row_cells[4].text = str(conArr2[i])
        row_cells[5].text = str(presArr1[i])
        row_cells[6].text = str(presArr2[i])

#~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
# add figure to report
    myReport.add_heading("Test Figure", 1)
    myReport.add_paragraph("Temperature Figure: ")
    myReport.add_picture("tempplot.jpg", width=docx.shared.Inches(4), height=docx.shared.Inches(4))
    myReport.add_paragraph("Conductivity Figure: ")
    myReport.add_picture("conplot.jpg", width=docx.shared.Inches(4), height=docx.shared.Inches(4))
    myReport.add_paragraph("Pressure Figure: ")
    myReport.add_picture("presplot.jpg", width=docx.shared.Inches(4), height=docx.shared.Inches(4))
#~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~

    myReport.add_heading("Test Result ", 1)
    myReport.add_paragraph("Note: If the difference of test CTD's data and standard CTD's data is within 1%, then the corresponding test is passed. If all the Temperature, Conductivity and Pressure(if exists) Tests passed, then final result will be passed.")

    final_pass_or_fail = False

    if (tempCum/myLen <= 0.01):
        myReport.add_paragraph("Temperature Test: Passed")
    else:
        myReport.add_paragraph("Temperature Test: Failed")

    if (conCum/myLen <= 0.01):
        myReport.add_paragraph("Conductivity Test: Passed")
    else:
        myReport.add_paragraph("Conductivity Test: Failed")

    
    if (isPres == 20):
        myReport.add_paragraph("Pressure Test: Passed")
    else:
        myReport.add_paragraph("Pressure Test: Failed")

    if (tempCum/myLen <= 0.01) and (conCum/myLen <= 0.01) and (isPres == 20):
        final_pass_or_fail = True
        final = myReport.add_paragraph("Final Result: ")
        final.add_run('Pass!').bold = True
    else:
        final = myReport.add_paragraph("Final Result: ")
        final.add_run('Fail!').bold = True
    
    #Title report in format "yyyy-mm-dd_SBE##_SN:xxxx_PASS" where ## = model(ie 37 or 19) and xxxx = serial number 
    if final_pass_or_fail:
        reportResult = "PASS"
    else:
        reportResult = "FAIL"

    reportTime = datetime.today().strftime('%Y-%m-%d')
    reportName = "%s_AMLMETRECX_SN%s_%s.docx" % (reportTime, test_sn, reportResult)

    try:
        myReport.save(r'%s' % reportName)
        # end words on GUI 
        ending = Label(mainWin, text = "Report successfully generated.", background = "#737CA1",  font=fontStyle, foreground = "white")
        ending.grid(row = 13, column = 0)
    except:
        ending = Label(mainWin, text = "Fail to generate the report.", background = "#737CA1",  font=fontStyle, foreground = "white")
        ending.grid(row = 13, column = 0)
        return

    tn1.close()
    tn2.close()
          
mainWin.mainloop()